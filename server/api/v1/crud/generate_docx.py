from typing import List
from io import BytesIO
from docx import Document

def generate_docx(data: List[dict]):
    buffer = BytesIO()
    doc = Document()
    for item in data:
        doc.add_paragraph(f"Question: {item['question']}")
        if 'options' in item:
            doc.add_paragraph("Options: " + ", ".join(item['options']))
        doc.add_paragraph(f"Answer: {item['answer']}")
        doc.add_paragraph("")  
    doc.save(buffer)
    buffer.seek(0)
    return buffer
