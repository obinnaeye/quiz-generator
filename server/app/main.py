from fastapi import FastAPI, Query, HTTPException
from app.api import healthcheck
from enum import Enum
from pydantic import BaseModel, EmailStr
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

from fastapi.responses import StreamingResponse
from io import StringIO, BytesIO
import csv
from typing import List
import json
from reportlab.pdfgen import canvas
from docx import Document

from reportlab.lib.pagesizes import letter

from .mock_quiz_data import quiz_data_multiple_choice, quiz_data_true_false, quiz_data_open_ended


app = FastAPI()


def generate_txt(data: List[dict]):
    buffer = StringIO()
    for item in data:
        buffer.write(f"Question: {item['question']}\n")
        if 'options' in item:
            buffer.write("Options: " + ", ".join(item['options']) + "\n")
        buffer.write(f"Answer: {item['answer']}\n\n")
    buffer.seek(0)
    return buffer

def generate_csv(data: List[dict]):
    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["Question", "Options", "Answer"])
    for item in data:
        options = ", ".join(item.get("options", []))
        writer.writerow([item["question"], options, item["answer"]])
    buffer.seek(0)
    return buffer

def generate_pdf(data: List[dict]):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    page_width, page_height = letter
    margin = 50  
    line_width = page_width - (2 * margin) 
    y_position = page_height - 50 

    for item in data:
        question = f"Question: {item['question']}"
        y_position = draw_wrapped_text(pdf, question, margin, y_position, line_width)

        if 'options' in item:
            options = "Options: " + ", ".join(item['options'])
            y_position = draw_wrapped_text(pdf, options, margin, y_position, line_width)

        answer = f"Answer: {item['answer']}"
        y_position = draw_wrapped_text(pdf, answer, margin, y_position, line_width)

        y_position -= 20

        if y_position < 50:
            pdf.showPage()
            y_position = page_height - 50

    pdf.save()
    buffer.seek(0)
    return buffer

def draw_wrapped_text(pdf, text, x, y, max_width, line_height=15):
    words = text.split()
    line = ""
    for word in words:
        if pdf.stringWidth(line + word + " ", "Helvetica", 12) <= max_width:
            line += word + " "
        else:
            pdf.drawString(x, y, line.strip())
            y -= line_height
            line = word + " "

    if line:
        pdf.drawString(x, y, line.strip())
        y -= line_height

    return y


def generate_docx(data: List[dict]):
    buffer = BytesIO()
    doc = Document()
    for item in data:
        doc.add_paragraph(f"Question: {item['question']}")
        if 'options' in item:
            doc.add_paragraph("Options: " + ", ".join(item['options']))
        doc.add_paragraph(f"Answer: {item['answer']}")
        doc.add_paragraph("")  
    doc.save(buffer)
    buffer.seek(0)
    return buffer


app.include_router(healthcheck.router, prefix="/api", tags=["healthcheck"])

@app.get("/api")
def read_root():
    return {"message": "Welcome to the Quiz App API!"}


class User(BaseModel):
    username: str
    email: EmailStr
    password: str

class LoginRequest(BaseModel):
    username_or_email: str
    password: str

class LoginResponse(BaseModel):
    message: str
    user: User

mock_db: list[User] = []


# The Register functionality
@app.post("/register/", response_model=User)
def create_user(user:User):
    if any(existing_user.username == user.username for existing_user in mock_db):
        raise HTTPException(status_code=400, detail="Username already taken")

    if any(existing_user.email == user.email for existing_user in mock_db):
        raise HTTPException(status_code=400, detail="Email already registered")
    
    mock_db.append(user)
    return user

# List all the users
@app.get("/users/", response_model=list[User])
def list_users():
    return mock_db



@app.post("/login/", response_model=LoginResponse)
def login(request: LoginRequest):  
    user = next((u for u in mock_db if (u.username == request.username_or_email or u.email == request.username_or_email) and u.password == request.password), None)

    if user is None:
        raise HTTPException(status_code=401, detail="Invalid credentials")

    return {"message": "Login successful", "user": user}

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Or use ["*"] to allow all origins (less secure)
    allow_credentials=True,
    allow_methods=["*"],  # Allow all HTTP methods (GET, POST, etc.)
    allow_headers=["*"],  # Allow all headers
)

# @app.get("/api/generate-quiz")
# def generate_quiz():
#     return {"message": "quiz generated"}




@app.get("/generate-quiz")
def generate_quiz(question_type: str = Query(..., description="Type of questions requested (multichoice, true-false, open-ended)")):
    if question_type == "multichoice":
        return {"quiz_data": quiz_data_multiple_choice}
    elif question_type == "true-false":
        return {"quiz_data": quiz_data_true_false}
    elif question_type == "open-ended":
        return {"quiz_data": quiz_data_open_ended}
    else:
        raise HTTPException(status_code=400, detail="Invalid question type")

@app.get("/download-quiz/")
async def download_quiz(format: str = "txt", type: str = "multichoice"):
    if type == "multichoice":
        quiz_data = quiz_data_multiple_choice
    elif type == "true-false":
        quiz_data = quiz_data_true_false
    elif type == "open-ended":
        quiz_data = quiz_data_open_ended
    else :
        raise HTTPException(status_code=400, detail="Unsupported question_type")

    
    if format == "txt":
        buffer = generate_txt(quiz_data)
        content_type = "text/plain"
    elif format == "csv":
        buffer = generate_csv(quiz_data)
        content_type = "text/csv"
    elif format == "pdf":
        buffer = generate_pdf(quiz_data)
        content_type = "application/pdf"
    elif format == "docx":
        buffer = generate_docx(quiz_data)
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format")

    return StreamingResponse(
        buffer, media_type=content_type, headers={
            "Content-Disposition": f"attachment; filename=quiz_data.{format}"
        }
    )
